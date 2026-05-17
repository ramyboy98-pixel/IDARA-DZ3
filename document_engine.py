import os
import re
import logging
from datetime import datetime
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

from utils.paths import get_output_dir

# إعداد النظام الداخلي للتسجيل
logger = logging.getLogger(__name__)

OUTPUT_FOLDER = get_output_dir()
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*(.*?)\s*\}\}")


def ensure_output_folder():
    """إنشاء مجلد الإخراج إذا لم يكن موجوداً"""
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def safe_filename(name):
    """تنظيف اسم الملف من الأحرف المحرمة
    
    Args:
        name: اسم الملف المراد تنظيفه
        
    Returns:
        str: اسم ملف آمن وصالح
    """
    value = str(name).strip()
    if not value:
        value = "document"

    forbidden = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in forbidden:
        value = value.replace(char, "_")

    value = value.replace(" ", "_")
    return value


def make_placeholder_key(label):
    """تحويل اسم الخانة إلى مفتاح موحد
    
    تحويل الأسماء العربية والإنجليزية إلى مفاتيح آمنة:
    - "تاريخ الطلب" => "تاريخ_الطلب"
    - "رقم الهاتف" => "رقم_الهاتف"
    
    الفائدة: يدعم الكود مع docx بدون استخدام docxtpl
    الذي يسبب مشاكل مع المتغيرات العربية
    
    Args:
        label: اسم الخانة/المتغير
        
    Returns:
        str: المفتاح الآمن والموحد
    """
    key = str(label or "").strip()
    replacements = [
        (" ", "_"), ("-", "_"), ("/", "_"), ("\\", "_"),
        (":", "_"), (";", "_"), (",", "_"), ("،", "_"),
        (".", "_"), ("(", ""), (")", ""), ("[", ""), ("]", ""),
    ]
    for old, new in replacements:
        key = key.replace(old, new)
    while "__" in key:
        key = key.replace("__", "_")
    return key.strip("_")


def normalize_data(data):
    """تطبيع البيانات للاستخدام مع المتغيرات
    
    تنشئ قاموس موحد يحتوي على كل المفاتيح الممكنة
    (الأصلية والآمنة) لكل قيمة
    
    Args:
        data: قاموس البيانات الأصلي
        
    Returns:
        dict: قاموس معياري مع جميع الصيغ الممكنة
    """
    normalized = {}
    for key, value in (data or {}).items():
        text_value = "" if value is None else str(value)
        raw_key = str(key or "").strip()
        safe_key = make_placeholder_key(raw_key)
        if raw_key:
            normalized[raw_key] = text_value
        if safe_key:
            normalized[safe_key] = text_value
    return normalized


def render_text_template(template_content, data):
    """استبدال المتغيرات في نص القالب
    
    يبحث عن متغيرات بصيغة {{اسم_المتغير}} ويستبدلها
    بقيمها من البيانات المعطاة
    
    Args:
        template_content: نص القالب يحتوي على متغيرات
        data: قاموس البيانات المراد استخدامها
        
    Returns:
        str: النص بعد استبدال جميع المتغيرات
        
    Example:
        >>> template = "اسم العميل: {{name}}"
        >>> data = {"name": "أحمد"}
        >>> render_text_template(template, data)
        'اسم العميل: أحمد'
    """
    text = template_content or ""
    values = normalize_data(data)

    def replace_match(match):
        original_key = match.group(1).strip()
        safe_key = make_placeholder_key(original_key)
        if original_key in values:
            return values[original_key]
        if safe_key in values:
            return values[safe_key]
        return match.group(0)

    return PLACEHOLDER_PATTERN.sub(replace_match, text)


def replace_text_in_paragraph(paragraph, data):
    """استبدال المتغيرات في فقرة من ملف Word
    
    يعمل على مستوى الـ runs الفردية دون تعديل البنية
    لحفظ التنسيق والمحاذاة والخصائص الأخرى
    
    السبب: تعديل paragraph.text مباشرة يفقد التنسيق
    الحل: العمل على مستوى الـ runs والتعديل الدقيق
    
    Args:
        paragraph: الفقرة من الوثيقة
        data: قاموس البيانات
    """
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text or "}}" not in full_text:
        return

    values = normalize_data(data)
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    def replacement_for(match):
        original_key = match.group(1).strip()
        safe_key = make_placeholder_key(original_key)
        if original_key in values:
            return values[original_key]
        if safe_key in values:
            return values[safe_key]
        return match.group(0)

    # خريطة كل حرف إلى رقم الـ run ومكانه داخله
    char_map = []
    for run_index, run in enumerate(paragraph.runs):
        for char_index, _ in enumerate(run.text):
            char_map.append((run_index, char_index))

    # معالجة من نهاية الفقرة إلى بدايتها لتجنب تغيير الإحداثيات
    for match in reversed(matches):
        replacement = replacement_for(match)
        if replacement == match.group(0):
            continue

        start = match.start()
        end = match.end() - 1
        if start >= len(char_map) or end >= len(char_map):
            continue

        start_run_index, start_char_index = char_map[start]
        end_run_index, end_char_index = char_map[end]

        if start_run_index == end_run_index:
            run = paragraph.runs[start_run_index]
            text = run.text
            run.text = text[:start_char_index] + replacement + text[end_char_index + 1:]
            continue

        start_run = paragraph.runs[start_run_index]
        end_run = paragraph.runs[end_run_index]

        start_prefix = start_run.text[:start_char_index]
        end_suffix = end_run.text[end_char_index + 1:]

        start_run.text = start_prefix + replacement
        for idx in range(start_run_index + 1, end_run_index):
            paragraph.runs[idx].text = ""
        end_run.text = end_suffix


def replace_placeholders_in_document(document, data):
    """استبدال جميع المتغيرات في كامل الوثيقة
    
    يعالج:
    - الفقرات في النص الرئيسي
    - الجداول وخلاياها
    - رأس وتذييل الصفحة
    
    Args:
        document: كائن الوثيقة من python-docx
        data: قاموس البيانات
    """
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, data)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)

    # بعض القوالب تضع المتغيرات داخل رأس أو تذييل الصفحة
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, data)


def generate_word_document(template_path, data, template_name="وثيقة"):
    """إنشاء ملف Word من قالب مرفوع
    
    يحمل القالب ويستبدل جميع المتغيرات ثم يحفظها
    
    ملاحظة مهمة:
    - لا نستخدم docxtpl لأنه يسبب مشاكل مع العربية
    - البرنامج يستبدل المتغيرات بنفسه بطريقة آمنة
    - يدعم الصيغ: {{تاريخ الطلب}} و {{تاريخ_الطلب}}
    
    Args:
        template_path: مسار ملف القالب Word
        data: قاموس البيانات
        template_name: اسم القالب (للملف الناتج)
        
    Returns:
        str: مسار الملف الناتج
        
    Raises:
        FileNotFoundError: إذا لم يوجد ملف القالب
        Exception: عند فشل معالجة الملف
    """
    ensure_output_folder()

    try:
        document = Document(template_path)
        logger.debug(f"تم تحميل قالب Word: {template_path}")
    except FileNotFoundError:
        logger.error(f"ملف القالب غير موجود: {template_path}")
        raise
    except Exception as e:
        logger.error(f"خطأ في تحميل قالب Word: {str(e)}")
        raise

    replace_placeholders_in_document(document, data)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_filename(template_name)}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    
    try:
        document.save(output_path)
        logger.info(f"تم إنشاء وثيقة Word: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"خطأ في حفظ وثيقة Word: {str(e)}")
        raise


def generate_word_from_text_template(template_content, data, template_name="وثيقة"):
    """إنشاء ملف Word من نص القالب
    
    يعالج النص ثم يضعه في وثيقة Word جديدة
    
    Args:
        template_content: محتوى القالب النصي
        data: قاموس البيانات
        template_name: اسم القالب
        
    Returns:
        str: مسار الملف الناتج
        
    Raises:
        Exception: عند فشل إنشاء الوثيقة
    """
    ensure_output_folder()
    rendered_text = render_text_template(template_content, data)

    document = Document()

    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(56)
    section.right_margin = Pt(56)

    for line in rendered_text.splitlines():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(14)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_filename(template_name)}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    
    try:
        document.save(output_path)
        logger.info(f"تم إنشاء وثيقة Word من نص: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"خطأ في حفظ وثيقة Word: {str(e)}")
        raise


def convert_word_to_pdf(word_path):
    """تحويل ملف Word إلى PDF باستخدام Microsoft Word
    
    السبب:
    إنشاء PDF يدويا عبر reportlab لا يحافظ على العربية والتنسيق
    يظهر نص غريب مثل ï؟½. الحل: استخدام Word نفسه للتصدير
    
    ملاحظة حرجة:
    - هذه الدالة تعمل على Windows فقط
    - يجب تثبيت Microsoft Word
    - تتطلب مكتبة pywin32
    
    Args:
        word_path: مسار ملف Word المراد تحويله
        
    Returns:
        str: مسار ملف PDF الناتج
        
    Raises:
        FileNotFoundError: إذا لم يوجد ملف Word
        RuntimeError: إذا لم يكن Windows أو Word غير مثبت
        ImportError: إذا كانت pywin32 غير مثبت��
    """
    if not word_path or not os.path.exists(word_path):
        logger.error(f"ملف وورد غير موجود: {word_path}")
        raise FileNotFoundError("ملف وورد غير موجود للتحويل إلى PDF")

    pdf_path = os.path.splitext(word_path)[0] + ".pdf"

    if os.name != "nt":
        logger.error("محاولة تحويل PDF على نظام غير Windows")
        raise RuntimeError("تحويل وورد إلى PDF يحتاج Windows و Microsoft Word")

    try:
        import win32com.client
    except ImportError:
        logger.error("مكتبة pywin32 غير مثبتة")
        raise RuntimeError("مكتبة pywin32 غير مثبتة. أضف pywin32 إلى requirements.txt ثم أعد البناء.")

    word_app = None
    document = None
    try:
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0

        abs_word_path = os.path.abspath(word_path)
        abs_pdf_path = os.path.abspath(pdf_path)

        document = word_app.Documents.Open(abs_word_path)
        # 17 = wdFormatPDF
        document.SaveAs(abs_pdf_path, FileFormat=17)
        document.Close(False)
        document = None
        
        logger.info(f"تم تحويل Word إلى PDF: {abs_pdf_path}")
        return abs_pdf_path

    except Exception as e:
        logger.error(f"خطأ في تحويل Word إلى PDF: {str(e)}")
        raise
    finally:
        try:
            if document is not None:
                document.Close(False)
        except Exception:
            pass
        try:
            if word_app is not None:
                word_app.Quit()
        except Exception:
            pass


def generate_simple_pdf(data, template_name="وثيقة", template_content=None):
    """إنشاء ملف PDF بسيط باستخدام reportlab
    
    ملاحظة:
    هذه الدالة تستخدم للملفات البسيطة فقط
    لا تحافظ تماماً على تنسيق Word المعقد
    
    Args:
        data: قاموس البيانات
        template_name: اسم المستند
        template_content: محتوى القالب (اختياري)
        
    Returns:
        str: مسار ملف PDF الناتج
        
    Raises:
        Exception: عند فشل إنشاء PDF
    """
    ensure_output_folder()

    pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_filename(template_name)}_{timestamp}.pdf"
    pdf_path = os.path.join(OUTPUT_FOLDER, filename)

    doc = SimpleDocTemplate(pdf_path)
    styles = getSampleStyleSheet()
    story = []

    title = Paragraph(
        f"<font name='HYSMyeongJo-Medium'><b>{escape(str(template_name))}</b></font>",
        styles['Title']
    )
    story.append(title)
    story.append(Spacer(1, 20))

    try:
        if template_content:
            rendered = render_text_template(template_content, data)
            for line in rendered.splitlines():
                text = f"<font name='HYSMyeongJo-Medium'>{escape(str(line))}</font>"
                story.append(Paragraph(text, styles['BodyText']))
                story.append(Spacer(1, 8))
        else:
            for key, value in (data or {}).items():
                text = f"<font name='HYSMyeongJo-Medium'>{escape(str(key))}: {escape(str(value))}</font>"
                story.append(Paragraph(text, styles['BodyText']))
                story.append(Spacer(1, 10))

        doc.build(story)
        logger.info(f"تم إنشاء ملف PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"خطأ في إنشاء ملف PDF: {str(e)}")
        raise
